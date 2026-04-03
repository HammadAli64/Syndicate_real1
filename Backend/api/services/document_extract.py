"""
Document text extraction for the syndicate pipeline.

Pipeline alignment:
  1) Upload → bytes in memory → UploadedDocument.text_extracted (optional: Cloudinary for raw files later).
  2) Extract text here (PDF via PyMuPDF when installed, else pypdf; DOCX via python-docx).
  3) OpenAI ingests text → MindsetKnowledge.payload (JSON: mindsets, themes, etc.).
  4) Challenges app reads DB and generates daily tasks / scoring (see apps.challenges).
"""
from __future__ import annotations

import hashlib
import io
import logging
from pathlib import Path

from django.conf import settings

logger = logging.getLogger(__name__)


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _pdf_text_pymupdf(data: bytes) -> str | None:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            parts: list[str] = []
            for page in doc:
                t = page.get_text() or ""
                if t.strip():
                    parts.append(t)
            return "\n\n".join(parts).strip()
        finally:
            doc.close()
    except Exception as e:
        logger.warning("PyMuPDF extract failed, will try pypdf: %s", e)
        return None


def _pdf_text_pymupdf_path(path: Path) -> str | None:
    try:
        import fitz
    except ImportError:
        return None
    try:
        doc = fitz.open(path)
        try:
            parts: list[str] = []
            for page in doc:
                t = page.get_text() or ""
                if t.strip():
                    parts.append(t)
            return "\n\n".join(parts).strip()
        finally:
            doc.close()
    except Exception as e:
        logger.warning("PyMuPDF path extract failed, will try pypdf: %s", e)
        return None


def _pdf_text_pypdf_bytes(data: bytes) -> str:
    from pypdf import PdfReader

    buf = io.BytesIO(data)
    reader = PdfReader(buf)
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip()


def _pdf_text_pypdf_path(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip()


def _pdf_text_bytes(data: bytes) -> str:
    text = _pdf_text_pymupdf(data)
    if text is not None and text.strip():
        return text
    return _pdf_text_pypdf_bytes(data)


def extract_text_from_bytes(data: bytes, suffix: str) -> str:
    """Parse upload bytes in memory (no temp file). suffix includes leading dot, e.g. .docx."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        return _pdf_text_bytes(data)
    if suffix in {".txt", ".md", ".markdown"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return _docx_text_bytes(data)
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, .md, or .docx")


def _docx_text_bytes(data: bytes) -> str:
    import docx

    buf = io.BytesIO(data)
    document = docx.Document(buf)
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _pdf_text_pymupdf_path(path)
        if text is not None and text.strip():
            return text
        return _pdf_text_pypdf_path(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _docx_text(path)
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, .md, or .docx")


def _docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _pdf_text(path: Path) -> str:
    text = _pdf_text_pymupdf_path(path)
    if text is not None and text.strip():
        return text
    return _pdf_text_pypdf_path(path)


def truncate(text: str, max_chars: int | None = None) -> str:
    max_chars = max_chars or getattr(settings, "SYNDICATE_MAX_DOC_CHARS", 120_000)
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[... document truncated for processing ...]"
